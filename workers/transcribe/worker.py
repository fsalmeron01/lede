import json
import os
import pathlib
import time
from datetime import datetime, timezone

from docx import Document
from faster_whisper import WhisperModel
import psycopg
from redis import Redis


TRANSCRIBE_LIST_KEY = "transcribe:jobs"

def env(name, default=None):
    return os.getenv(name, default)

DATABASE_URL = env("DATABASE_URL")
REDIS_URL = env("REDIS_URL", "redis://redis:6379")
STORAGE_ROOT = env("STORAGE_ROOT", "/data")
WHISPER_MODEL = env("WHISPER_MODEL", "base")
WHISPER_DEVICE = env("WHISPER_DEVICE", "cpu")

_model = None

def get_model():
    global _model
    if _model is None:
        print(f"[transcribe] Loading Whisper model: {WHISPER_MODEL} on {WHISPER_DEVICE}", flush=True)
        _model = WhisperModel(WHISPER_MODEL, device=WHISPER_DEVICE)
        print(f"[transcribe] Model loaded.", flush=True)
    return _model

def get_pg_conn():
    return psycopg.connect(DATABASE_URL)

def get_redis():
    return Redis.from_url(REDIS_URL, decode_responses=True)

def format_timestamp(seconds):
    total = int(seconds)
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02d}:{m:02d}:{s:02d}"

def update_job(job_id, **patch):
    if not patch:
        return
    assignments = []
    values = []
    for key, value in patch.items():
        assignments.append(f"{key} = %s")
        values.append(value)
    values.append(job_id)
    sql = f"UPDATE jobs SET {', '.join(assignments)} WHERE id = %s"
    with get_pg_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, values)
        conn.commit()

def get_job(job_id):
    with get_pg_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, source_url, source_type, title, requested_outputs, content_mode, metadata FROM jobs WHERE id = %s",
                (job_id,),
            )
            row = cur.fetchone()
            if not row:
                return None
            return {
                "id": row[0],
                "source_url": row[1],
                "source_type": row[2],
                "title": row[3],
                "requested_outputs": row[4] or [],
                "content_mode": row[5] or "camden-tribune",
                "metadata": row[6] or {},
            }

def insert_artifact(job_id, artifact_type, file_path, mime_type):
    stat = os.stat(file_path)
    with get_pg_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO artifacts (job_id, artifact_type, file_path, file_name, mime_type, size_bytes) VALUES (%s, %s, %s, %s, %s, %s)",
                (job_id, artifact_type, file_path, os.path.basename(file_path), mime_type, stat.st_size),
            )
        conn.commit()

def upsert_transcript(job_id, language, raw_text, clean_text, segments_json):
    with get_pg_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO transcripts (job_id, language, raw_text, clean_text, segments_json)
                VALUES (%s, %s, %s, %s, %s::jsonb)
                ON CONFLICT (job_id) DO UPDATE SET
                    language = EXCLUDED.language,
                    raw_text = EXCLUDED.raw_text,
                    clean_text = EXCLUDED.clean_text,
                    segments_json = EXCLUDED.segments_json
                """,
                (job_id, language, raw_text, clean_text, json.dumps(segments_json)),
            )
        conn.commit()

def build_txt(job, output_path, clean_text):
    lines = [
        "Transcript Export",
        "=================",
        f"Job ID: {job['id']}",
        f"Source: {job['source_url']}",
        f"Title: {job.get('title') or 'N/A'}",
        f"Content Mode: {job.get('content_mode', 'camden-tribune')}",
        f"Processed: {datetime.now(timezone.utc).isoformat()}",
        "",
        "Transcript",
        "----------",
        clean_text,
        "",
    ]
    pathlib.Path(output_path).write_text("\n".join(lines), encoding="utf-8")

def build_docx(job, output_path, clean_text, segments_json):
    doc = Document()
    doc.add_heading("Transcript Export", level=1)
    doc.add_paragraph(f"Job ID: {job['id']}")
    doc.add_paragraph(f"Source: {job['source_url']}")
    doc.add_paragraph(f"Title: {job.get('title') or 'N/A'}")
    doc.add_paragraph(f"Content Mode: {job.get('content_mode', 'camden-tribune')}")
    doc.add_paragraph(f"Processed: {datetime.now(timezone.utc).isoformat()}")
    doc.add_heading("Transcript", level=2)
    doc.add_paragraph(clean_text or "")
    doc.add_heading("Timestamped Segments", level=2)
    for seg in segments_json:
        start = seg.get("start", 0.0)
        end = seg.get("end", 0.0)
        text = seg.get("text", "").strip()
        doc.add_paragraph(f"[{format_timestamp(start)} → {format_timestamp(end)}] {text}")
    doc.save(output_path)

def transcribe(job_id, mp3_path):
    job = get_job(job_id)
    if not job:
        raise RuntimeError(f"Job {job_id} not found.")

    print(f"[transcribe] Job {job_id} — starting transcription", flush=True)
    print(f"[transcribe] File: {mp3_path}", flush=True)

    # Check file exists and size
    if not os.path.exists(mp3_path):
        raise RuntimeError(f"Audio file not found: {mp3_path}")

    file_size_mb = os.path.getsize(mp3_path) / (1024 * 1024)
    print(f"[transcribe] Audio file size: {file_size_mb:.1f} MB", flush=True)

    update_job(job_id, status="transcribing", progress=65)

    model = get_model()

    print(f"[transcribe] Job {job_id} — running Whisper ({WHISPER_MODEL}/{WHISPER_DEVICE})", flush=True)
    start_time = time.time()

    segments_iter, info = model.transcribe(mp3_path, vad_filter=True)

    print(f"[transcribe] Detected language: {getattr(info, 'language', 'unknown')} "
          f"(prob: {getattr(info, 'language_probability', 0):.2f})", flush=True)
    print(f"[transcribe] Duration: {getattr(info, 'duration', 0):.1f}s", flush=True)

    raw_parts = []
    clean_parts = []
    segment_rows = []
    segment_count = 0
    last_progress_update = time.time()

    for segment in segments_iter:
        text = (segment.text or "").strip()
        raw_parts.append(f"[{format_timestamp(segment.start)} - {format_timestamp(segment.end)}] {text}")
        if text:
            clean_parts.append(text)
        segment_rows.append({"start": segment.start, "end": segment.end, "text": text})
        segment_count += 1

        # Log progress every 60 seconds and update DB to prevent "stuck" detection
        now = time.time()
        if now - last_progress_update > 60:
            duration = getattr(info, 'duration', 0)
            if duration > 0:
                pct = min(int(65 + (segment.end / duration) * 20), 84)
                elapsed = now - start_time
                print(f"[transcribe] Progress: {segment.end:.0f}s / {duration:.0f}s "
                      f"({pct}%) — {segment_count} segments — elapsed: {elapsed:.0f}s", flush=True)
                update_job(job_id, status="transcribing", progress=pct)
            last_progress_update = now

    elapsed = time.time() - start_time
    print(f"[transcribe] Job {job_id} — transcription complete: "
          f"{segment_count} segments in {elapsed:.1f}s", flush=True)

    raw_text = "\n".join(raw_parts).strip()
    clean_text = "\n".join(clean_parts).strip()

    upsert_transcript(
        job_id=job_id,
        language=getattr(info, "language", None),
        raw_text=raw_text,
        clean_text=clean_text,
        segments_json=segment_rows,
    )

    output_dir = pathlib.Path(STORAGE_ROOT) / "jobs" / job_id / "downloads"
    output_dir.mkdir(parents=True, exist_ok=True)

    requested = set(job.get("requested_outputs") or [])

    if "txt" in requested:
        txt_path = output_dir / f"{job_id}.txt"
        build_txt(job, str(txt_path), clean_text)
        insert_artifact(job_id, "txt", str(txt_path), "text/plain")
        print(f"[transcribe] TXT artifact saved", flush=True)

    if "docx" in requested:
        docx_path = output_dir / f"{job_id}.docx"
        build_docx(job, str(docx_path), clean_text, segment_rows)
        insert_artifact(job_id, "docx", str(docx_path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        print(f"[transcribe] DOCX artifact saved", flush=True)

    json_path = output_dir / f"{job_id}.segments.json"
    json_path.write_text(json.dumps(segment_rows, ensure_ascii=False, indent=2), encoding="utf-8")
    insert_artifact(job_id, "json", str(json_path), "application/json")

    # Signal LLM worker
    update_job(job_id, status="generating-article", progress=88)
    print(f"[transcribe] Job {job_id} — complete, signaling LLM worker", flush=True)

def main():
    redis = get_redis()
    print(f"[transcribe] Worker ready — model: {WHISPER_MODEL} device: {WHISPER_DEVICE}", flush=True)

    while True:
        item = redis.brpop(TRANSCRIBE_LIST_KEY, timeout=5)
        if not item:
            continue
        _, raw = item
        job_id = None
        try:
            payload = json.loads(raw)
            job_id = payload.get("jobId")
            mp3_path = payload.get("mp3Path")
            if not job_id or not mp3_path:
                print(f"[transcribe] Invalid payload: {raw}", flush=True)
                continue
            print(f"[transcribe] Received job {job_id}", flush=True)
            transcribe(job_id, mp3_path)
        except Exception as exc:
            print(f"[transcribe] ERROR job {job_id}: {exc}", flush=True)
            import traceback
            traceback.print_exc()
            if job_id:
                update_job(job_id, status="failed", progress=0, error_message=str(exc))

if __name__ == "__main__":
    main()
